import re
from io import BytesIO
from pathlib import Path
from zipfile import ZIP_DEFLATED, ZipFile

import pytest
from django.core.exceptions import ValidationError
from django.core.files.uploadedfile import SimpleUploadedFile
from django.urls import reverse
from django.utils import timezone
from docx import Document
from rest_framework.test import APIClient

from apps.blog.importers import NoteImportError, extract_uploaded_note
from apps.blog.models import Article, ArticleSourceFile, Category


@pytest.fixture
def api_client():
    return APIClient()


@pytest.fixture
def note_upload_root(settings):
    directory = Path(settings.REPOSITORY_DIR) / "data" / "notes"
    directory.mkdir(parents=True, exist_ok=True)
    existing_files = {path.resolve() for path in directory.rglob("*") if path.is_file()}
    yield directory
    for path in directory.rglob("*"):
        if path.is_file() and path.resolve() not in existing_files:
            path.unlink()


def make_pdf(text: str = "Readable PDF note") -> bytes:
    content = f"BT /F1 12 Tf 72 720 Td ({text}) Tj ET".encode()
    objects = [
        b"<< /Type /Catalog /Pages 2 0 R >>",
        b"<< /Type /Pages /Kids [3 0 R] /Count 1 >>",
        b"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] "
        b"/Resources << /Font << /F1 5 0 R >> >> /Contents 4 0 R >>",
        b"<< /Length " + str(len(content)).encode() + b" >>\nstream\n" + content + b"\nendstream",
        b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>",
    ]
    payload = bytearray(b"%PDF-1.4\n")
    offsets = [0]
    for index, obj in enumerate(objects, start=1):
        offsets.append(len(payload))
        payload.extend(f"{index} 0 obj\n".encode())
        payload.extend(obj)
        payload.extend(b"\nendobj\n")
    xref_offset = len(payload)
    payload.extend(f"xref\n0 {len(objects) + 1}\n".encode())
    payload.extend(b"0000000000 65535 f \n")
    for offset in offsets[1:]:
        payload.extend(f"{offset:010d} 00000 n \n".encode())
    payload.extend(
        f"trailer\n<< /Size {len(objects) + 1} /Root 1 0 R >>\n"
        f"startxref\n{xref_offset}\n%%EOF\n".encode()
    )
    return bytes(payload)


def make_docx() -> bytes:
    document = Document()
    document.add_heading("Word 导入标题", level=1)
    document.add_paragraph("这是从 DOCX 提取的正文。")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "字段"
    table.cell(0, 1).text = "说明"
    table.cell(1, 0).text = "格式"
    table.cell(1, 1).text = "DOCX"
    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def test_markdown_docx_and_pdf_extract_real_text(settings):
    settings.NOTE_UPLOAD_MAX_BYTES = 8 * 1024 * 1024
    markdown = extract_uploaded_note(
        SimpleUploadedFile(
            "note.md",
            "# Markdown 标题\n\n正文内容".encode(),
            content_type="text/markdown",
        )
    )
    docx = extract_uploaded_note(
        SimpleUploadedFile(
            "note.docx",
            make_docx(),
            content_type=(
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            ),
        )
    )
    pdf = extract_uploaded_note(
        SimpleUploadedFile("note.pdf", make_pdf(), content_type="application/pdf")
    )

    assert markdown.outline[0]["text"] == "Markdown 标题"
    assert "# Word 导入标题" in docx.body_markdown
    assert "| 字段 | 说明 |" in docx.body_markdown
    assert "Readable PDF note" in pdf.body_markdown

    scanned_pdf = extract_uploaded_note(
        SimpleUploadedFile("scan.pdf", make_pdf(""), content_type="application/pdf")
    )
    assert "原版视图" in scanned_pdf.body_markdown


def test_import_rejects_spoofed_and_unsafe_files(settings):
    settings.NOTE_UPLOAD_MAX_BYTES = 32
    with pytest.raises(NoteImportError, match="不能超过"):
        extract_uploaded_note(SimpleUploadedFile("large.md", b"x" * 64, content_type="text/plain"))

    settings.NOTE_UPLOAD_MAX_BYTES = 1024 * 1024
    with pytest.raises(NoteImportError, match="内容与扩展名"):
        extract_uploaded_note(
            SimpleUploadedFile("fake.pdf", b"not a pdf", content_type="application/pdf")
        )

    archive_buffer = BytesIO()
    with ZipFile(archive_buffer, "w", ZIP_DEFLATED) as archive:
        archive.writestr("[Content_Types].xml", "<Types />")
        archive.writestr("word/document.xml", "<document />")
        archive.writestr("../escape.txt", "unsafe")
    with pytest.raises(NoteImportError, match="不安全"):
        extract_uploaded_note(
            SimpleUploadedFile(
                "unsafe.docx",
                archive_buffer.getvalue(),
                content_type=(
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                ),
            )
        )


@pytest.mark.django_db
def test_category_rejects_ancestor_cycle():
    root = Category.objects.create(name="测试根分类", slug="test-root")
    child = Category.objects.create(name="测试子分类", slug="test-child", parent=root)
    root.parent = child

    with pytest.raises(ValidationError, match="循环"):
        root.full_clean()


@pytest.mark.django_db
def test_trusted_local_frontend_imports_and_downloads_markdown(
    api_client, category, settings, note_upload_root
):
    settings.DEBUG = True
    settings.NOTE_BROWSER_IMPORT_ENABLED = True
    settings.NOTE_UPLOAD_ROOT = note_upload_root
    settings.NOTE_UPLOAD_MAX_BYTES = 8 * 1024 * 1024
    settings.CORS_ALLOWED_ORIGINS = ["http://localhost:3000"]
    upload = SimpleUploadedFile(
        "../../research-note.md",
        "# 本地笔记\n\n安全保存并展示。".encode(),
        content_type="text/markdown",
    )

    response = api_client.post(
        reverse("article-import-file"),
        {"file": upload, "category_slug": category.slug},
        format="multipart",
        HTTP_ORIGIN="http://localhost:3000",
        REMOTE_ADDR="127.0.0.1",
    )

    assert response.status_code == 201
    assert response.data["title"] == "本地笔记"
    assert response.data["body_markdown"].startswith("# 本地笔记")
    article = Article.objects.get(slug=response.data["slug"])
    source = ArticleSourceFile.objects.get(article=article)
    stored_path = Path(source.file.path).resolve()
    assert stored_path.is_relative_to(note_upload_root.resolve())
    assert source.original_filename == "research-note.md"
    assert stored_path.name != source.original_filename

    detail = api_client.get(reverse("article-detail", kwargs={"slug": article.slug}))
    assert detail.status_code == 200
    assert set(detail.data["source_file"]) == {
        "original_filename",
        "source_format",
        "source_format_label",
        "size_bytes",
        "download_url",
        "preview_url",
    }
    assert detail.data["source_file"]["preview_url"] is None
    download = api_client.get(reverse("article-source-file", kwargs={"slug": article.slug}))
    assert download.status_code == 200
    assert download["X-Content-Type-Options"] == "nosniff"
    assert download["Content-Disposition"].startswith("attachment;")
    download.close()


@pytest.mark.django_db
def test_uploaded_pdf_has_inline_original_preview(
    api_client, category, settings, note_upload_root
):
    settings.DEBUG = True
    settings.NOTE_BROWSER_IMPORT_ENABLED = True
    settings.NOTE_UPLOAD_ROOT = note_upload_root
    settings.NOTE_UPLOAD_MAX_BYTES = 8 * 1024 * 1024
    settings.CORS_ALLOWED_ORIGINS = ["http://localhost:3000"]
    response = api_client.post(
        reverse("article-import-file"),
        {
            "file": SimpleUploadedFile(
                "original.pdf", make_pdf("Original layout"), content_type="application/pdf"
            ),
            "category_slug": category.slug,
        },
        format="multipart",
        HTTP_ORIGIN="http://localhost:3000",
        REMOTE_ADDR="127.0.0.1",
    )

    assert response.status_code == 201
    assert response.data["source_file"]["preview_url"].endswith("/preview-file/")
    preview = api_client.get(
        reverse("article-preview-file", kwargs={"slug": response.data["slug"]})
    )
    assert preview.status_code == 200
    assert preview["Content-Type"] == "application/pdf"
    assert preview["Content-Disposition"].startswith("inline;")
    assert preview["X-Content-Type-Options"] == "nosniff"
    preview.close()


@pytest.mark.django_db
def test_note_import_rejects_duplicate_and_untrusted_requests(
    api_client, category, settings, note_upload_root
):
    settings.DEBUG = True
    settings.NOTE_BROWSER_IMPORT_ENABLED = True
    settings.NOTE_UPLOAD_ROOT = note_upload_root
    settings.NOTE_UPLOAD_MAX_BYTES = 8 * 1024 * 1024
    settings.CORS_ALLOWED_ORIGINS = ["http://localhost:3000"]
    endpoint = reverse("article-import-file")
    content = b"# Duplicate\n\nBody"

    untrusted = api_client.post(
        endpoint,
        {
            "file": SimpleUploadedFile("note.md", content, content_type="text/plain"),
            "category_slug": category.slug,
        },
        format="multipart",
        HTTP_ORIGIN="https://malicious.example",
        REMOTE_ADDR="127.0.0.1",
    )
    assert untrusted.status_code == 403
    assert ArticleSourceFile.objects.count() == 0

    first = api_client.post(
        endpoint,
        {
            "file": SimpleUploadedFile("note.md", content, content_type="text/plain"),
            "category_slug": category.slug,
        },
        format="multipart",
        HTTP_ORIGIN="http://localhost:3000",
        REMOTE_ADDR="127.0.0.1",
    )
    duplicate = api_client.post(
        endpoint,
        {
            "file": SimpleUploadedFile("again.md", content, content_type="text/plain"),
            "category_slug": category.slug,
        },
        format="multipart",
        HTTP_ORIGIN="http://localhost:3000",
        REMOTE_ADDR="127.0.0.1",
    )
    assert first.status_code == 201
    assert duplicate.status_code == 409
    assert ArticleSourceFile.objects.count() == 1


@pytest.mark.django_db
def test_note_tree_returns_nested_metadata(api_client, published_article):
    response = api_client.get(reverse("article-tree"))

    assert response.status_code == 200
    categories = {item["slug"]: item for item in response.data["categories"]}
    assert categories["notes"]["parent_slug"] == "inference-optimization"
    assert [item["slug"] for item in categories["notes"]["ancestors"]] == [
        "ai-technology",
        "large-models",
        "inference-optimization",
    ]
    assert response.data["articles"][0]["slug"] == published_article.slug
    assert response.data["max_category_depth"] == 8


@pytest.mark.django_db
def test_trusted_local_frontend_creates_nested_note_directories(api_client, settings):
    settings.DEBUG = True
    settings.NOTE_BROWSER_IMPORT_ENABLED = True
    settings.CORS_ALLOWED_ORIGINS = ["http://localhost:3000"]
    endpoint = reverse("article-categories")
    request_meta = {
        "HTTP_ORIGIN": "http://localhost:3000",
        "REMOTE_ADDR": "127.0.0.1",
    }

    root_response = api_client.post(
        endpoint,
        {"name": "我的研究目录", "parent_slug": None},
        format="json",
        **request_meta,
    )
    assert root_response.status_code == 201
    assert root_response.data["parent_slug"] is None
    assert re.fullmatch(r"[a-z0-9-]+", root_response.data["slug"])
    assert len(root_response.data["slug"]) <= 100

    child_response = api_client.post(
        endpoint,
        {"name": "实验记录", "parent_slug": root_response.data["slug"]},
        format="json",
        **request_meta,
    )
    assert child_response.status_code == 201
    assert child_response.data["parent_slug"] == root_response.data["slug"]
    assert [item["slug"] for item in child_response.data["ancestors"]] == [
        root_response.data["slug"]
    ]
    assert (
        Category.objects.get(slug=child_response.data["slug"]).parent.slug
        == (root_response.data["slug"])
    )

    duplicate_response = api_client.post(
        endpoint,
        {"name": "实验记录", "parent_slug": None},
        format="json",
        **request_meta,
    )
    unsafe_response = api_client.post(
        endpoint,
        {"name": "../逃逸", "parent_slug": None},
        format="json",
        **request_meta,
    )
    assert duplicate_response.status_code == 400
    assert unsafe_response.status_code == 400


@pytest.mark.django_db
def test_note_directory_creation_rejects_excess_depth_and_untrusted_origin(api_client, settings):
    settings.DEBUG = True
    settings.NOTE_BROWSER_IMPORT_ENABLED = True
    settings.CORS_ALLOWED_ORIGINS = ["http://localhost:3000"]
    endpoint = reverse("article-categories")
    parent = Category.objects.create(name="层级 1", slug="level-1")
    for level in range(2, 9):
        parent = Category.objects.create(
            name=f"层级 {level}",
            slug=f"level-{level}",
            parent=parent,
        )

    too_deep = api_client.post(
        endpoint,
        {"name": "层级 9", "parent_slug": parent.slug},
        format="json",
        HTTP_ORIGIN="http://localhost:3000",
        REMOTE_ADDR="127.0.0.1",
    )
    untrusted = api_client.post(
        endpoint,
        {"name": "不受信目录", "parent_slug": None},
        format="json",
        HTTP_ORIGIN="https://malicious.example",
        REMOTE_ADDR="127.0.0.1",
    )

    assert too_deep.status_code == 400
    assert "最多支持 8 层" in str(too_deep.data)
    assert untrusted.status_code == 403
    assert not Category.objects.filter(name="不受信目录").exists()


@pytest.mark.django_db
def test_trusted_local_frontend_deletes_articles_and_empty_directories(
    api_client, django_capture_on_commit_callbacks, note_upload_root, settings
):
    settings.DEBUG = True
    settings.NOTE_BROWSER_IMPORT_ENABLED = True
    settings.NOTE_UPLOAD_ROOT = note_upload_root
    settings.CORS_ALLOWED_ORIGINS = ["http://localhost:3000"]
    request_meta = {
        "HTTP_ORIGIN": "http://localhost:3000",
        "REMOTE_ADDR": "127.0.0.1",
    }
    root = Category.objects.create(name="待删除根目录", slug="delete-root")
    child = Category.objects.create(name="待删除子目录", slug="delete-child", parent=root)
    article = Article.objects.create(
        title="待删除文章",
        slug="delete-article",
        summary="待删除",
        body_markdown="# 待删除",
        category=child,
        status=Article.Status.PUBLISHED,
        published_at=timezone.now(),
    )
    source = ArticleSourceFile.objects.create(
        article=article,
        file=SimpleUploadedFile("delete.md", b"# delete", content_type="text/markdown"),
        original_filename="delete.md",
        source_format=ArticleSourceFile.SourceFormat.MARKDOWN,
        content_type="text/markdown",
        size_bytes=8,
        sha256="d" * 64,
    )
    stored_path = Path(source.file.path)
    assert stored_path.exists()

    root_blocked = api_client.delete(
        reverse("article-category-detail", kwargs={"category_slug": root.slug}),
        **request_meta,
    )
    child_blocked = api_client.delete(
        reverse("article-category-detail", kwargs={"category_slug": child.slug}),
        **request_meta,
    )
    with django_capture_on_commit_callbacks(execute=True):
        article_deleted = api_client.delete(
            reverse("article-manage", kwargs={"slug": article.slug}),
            **request_meta,
        )
    child_deleted = api_client.delete(
        reverse("article-category-detail", kwargs={"category_slug": child.slug}),
        **request_meta,
    )
    root_deleted = api_client.delete(
        reverse("article-category-detail", kwargs={"category_slug": root.slug}),
        **request_meta,
    )

    assert root_blocked.status_code == 409
    assert "子目录" in root_blocked.data["detail"]
    assert child_blocked.status_code == 409
    assert "文章" in child_blocked.data["detail"]
    assert article_deleted.status_code == 204
    assert child_deleted.status_code == 204
    assert root_deleted.status_code == 204
    assert not Article.objects.filter(pk=article.pk).exists()
    assert not ArticleSourceFile.objects.filter(pk=source.pk).exists()
    assert not stored_path.exists()
    assert not Category.objects.filter(pk__in=(root.pk, child.pk)).exists()


@pytest.mark.django_db
def test_note_management_delete_rejects_untrusted_origin(api_client, published_article, settings):
    settings.DEBUG = True
    settings.NOTE_BROWSER_IMPORT_ENABLED = True
    settings.CORS_ALLOWED_ORIGINS = ["http://localhost:3000"]

    response = api_client.delete(
        reverse("article-manage", kwargs={"slug": published_article.slug}),
        HTTP_ORIGIN="https://malicious.example",
        REMOTE_ADDR="127.0.0.1",
    )

    assert response.status_code == 403
    assert Article.objects.filter(pk=published_article.pk).exists()
